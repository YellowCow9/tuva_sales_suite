"""
generate_exec_overview.py
--------------------------
Standalone script — run from tuva_tool/ directory:
    python generate_exec_overview.py

Requires: pip install python-docx
Output:   output/Tuva_Grant_Radar_Executive_Overview.docx
"""

import os
import pathlib
import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = pathlib.Path(__file__).parent
DATA_PATH  = SCRIPT_DIR / "data" / "recently_funded_profs" / "scored_leads.csv"
OUTPUT_DIR = SCRIPT_DIR / "output"
OUTPUT_PATH = OUTPUT_DIR / "Tuva_Grant_Radar_Executive_Overview.docx"

OUTPUT_DIR.mkdir(exist_ok=True)

TODAY = datetime.date.today().strftime("%B %d, %Y")

# ── Colors ────────────────────────────────────────────────────────────────────
TUVA_BLUE   = RGBColor(0x1F, 0x47, 0x88)   # #1F4788
TUVA_LIGHT  = RGBColor(0xD6, 0xE4, 0xF0)   # #D6E4F0 — header fill
HEADER_GRAY = RGBColor(0x40, 0x40, 0x40)   # dark gray text
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
SHADE_LIGHT = "D6E4F0"                      # hex for XML shading


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set cell background color via XML (python-docx has no direct API)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_color="1F4788", border_sz="4"):
    """Add a thin border around a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    border_sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = TUVA_BLUE
    return p


def add_body(doc: Document, text: str, italic: bool = False, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    return p


def add_callout(doc: Document, text: str):
    """Add a shaded callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EBF3FB")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size   = Pt(9.5)
    run.font.italic = True
    return p


# ── Load data ─────────────────────────────────────────────────────────────────

def load_top_leads(n: int = 4) -> list[dict]:
    df = pd.read_csv(DATA_PATH)
    top = df.sort_values("outbound_score", ascending=False).head(n)
    rows = []
    for _, r in top.iterrows():
        try:
            amt = f"${int(r['award_amount']):,}"
        except Exception:
            amt = str(r.get("award_amount", "N/A"))
        try:
            date_str = str(r["award_date"])[:7]   # "YYYY-MM"
        except Exception:
            date_str = str(r.get("award_date", "N/A"))
        rows.append({
            "tier":         str(r.get("tier", "—")),
            "pi_name":      str(r.get("pi_name", "—")).title(),
            "organization": str(r.get("organization", "—")).title()[:35],
            "award_amount": amt,
            "award_date":   date_str,
            "award_type":   str(r.get("award_type", "New")),
            "score":        f"{float(r.get('outbound_score', 0)):.3f}",
        })
    return rows


# ── Document builder ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover header ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Tuva Grant Radar — Competitive Intelligence Platform")
    title_run.bold = True
    title_run.font.size  = Pt(16)
    title_run.font.color.rgb = TUVA_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Executive Overview · {TODAY}")
    sub_run.italic = True
    sub_run.font.size  = Pt(11)
    sub_run.font.color.rgb = HEADER_GRAY
    sub_p.paragraph_format.space_after = Pt(14)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)  # spacer

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 1 — Lead Radar sample output
    # ═════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1 · Newly Funded PIs: Sample Output")
    add_body(doc,
        "The Lead Radar automatically identifies R01 recipients funded within the last 90 days "
        "and ranks them by fit for Tuva's value proposition. The table below shows the top 4 "
        "leads from the current dataset.")

    leads = load_top_leads(4)
    headers = ["Tier", "PI Name", "Institution", "Award Amount", "Award Date", "Type", "Score"]

    tbl = doc.add_table(rows=1 + len(leads), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], SHADE_LIGHT)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = TUVA_BLUE
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, lead in enumerate(leads, start=1):
        cells = tbl.rows[row_idx].cells
        values = [
            lead["tier"], lead["pi_name"], lead["organization"],
            lead["award_amount"], lead["award_date"], lead["award_type"], lead["score"],
        ]
        for col_idx, val in enumerate(values):
            cells[col_idx].text = val
            run = cells[col_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout(doc,
        "Score reflects a weighted composite of AI/data signal strength, grant freshness, "
        "award size, and award type across 8,700+ leads this quarter.")

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Grant Advisor sample output
    # ═════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2 · Grant Matching System: Sample Output")
    add_body(doc,
        "The Grant Advisor surfaces the 5 best-fit open funding opportunities for any PI based "
        "on semantic similarity to their full grant history, reranked by Claude AI.")

    grant_cards = [
        {
            "rank":  1,
            "score": "0.81",
            "title": "NIH — Advancing Multi-Omics Data Integration and Analysis",
            "close": "09/15/2026",
            "tags":  "Requires AI Methodology: Yes · Requires Data Plan: Yes",
            "blurb": (
                "Your prior work on single-cell RNA-seq and spatial transcriptomics aligns "
                "directly with this opportunity's emphasis on scalable computational frameworks "
                "for multi-omics data."
            ),
        },
        {
            "rank":  2,
            "score": "0.74",
            "title": "NIH — Biomedical Data Science Innovation Program",
            "close": "06/30/2026",
            "tags":  "",
            "blurb": (
                "Strong thematic overlap between your NIH-funded data harmonization pipeline "
                "and this grant's focus on reproducible, FAIR-compliant research infrastructure."
            ),
        },
    ]

    for card in grant_cards:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after  = Pt(2)
        run_h = p_title.add_run(f"Match #{card['rank']} · Score: {card['score']}")
        run_h.bold = True
        run_h.font.size = Pt(10)
        run_h.font.color.rgb = TUVA_BLUE

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        run_s = p_sub.add_run(card["title"])
        run_s.italic = True
        run_s.font.size = Pt(9.5)

        if card["tags"]:
            p_tags = doc.add_paragraph()
            p_tags.paragraph_format.space_after = Pt(2)
            run_t = p_tags.add_run(f"Close Date: {card['close']} · {card['tags']}")
            run_t.font.size = Pt(9)
            run_t.font.color.rgb = HEADER_GRAY
        else:
            p_tags = doc.add_paragraph()
            p_tags.paragraph_format.space_after = Pt(2)
            run_t = p_tags.add_run(f"Close Date: {card['close']}")
            run_t.font.size = Pt(9)
            run_t.font.color.rgb = HEADER_GRAY

        add_callout(doc, f'"{card["blurb"]}"')

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Projected value
    # ═════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3 · Projected Value for Tuva")

    metrics = [
        ("Leads identified per quarter",  "8,700+ NIH R01 recipients"),
        ("Tier 1 (highest-fit) leads",    "~870 per quarter (top 10%)"),
        ("Email coverage",                "~85% (3-stage inference pipeline)"),
        ("Grant freshness window",        "90 days from award — active budget decisions"),
        ("Manual research replaced",      "~1,450 hrs/qtr (10 min/lead × 8,700)"),
    ]

    mtbl = doc.add_table(rows=1 + len(metrics), cols=2)
    mtbl.style = "Table Grid"
    mtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    mhdr = mtbl.rows[0].cells
    for col, label in enumerate(["Metric", "Value"]):
        mhdr[col].text = label
        set_cell_bg(mhdr[col], SHADE_LIGHT)
        run = mhdr[col].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = TUVA_BLUE

    for row_idx, (metric, value) in enumerate(metrics, start=1):
        cells = mtbl.rows[row_idx].cells
        cells[0].text = metric
        cells[1].text = value
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    bullets = [
        ("Timing advantage",
         "Outreach reaches PIs while new-grant budgets are unallocated and infrastructure "
         "decisions are still being made — the highest-leverage window."),
        ("Fit signal",
         "Each lead includes abstract-level AI/data keyword analysis, meaning reps contact "
         "researchers whose work already matches Tuva's value proposition."),
        ("Grant Advisor as retention driver",
         "Existing Tuva clients get personalized grant recommendations, turning the platform "
         "into a strategic research partner."),
        ("Integration-ready",
         "JSON export maps directly to Clay, HubSpot, and Smartlead for immediate sequencing "
         "— no manual data entry."),
    ]

    for bold_label, body_text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run_b = p.add_run(f"{bold_label}: ")
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_body = p.add_run(body_text)
        run_body.font.size = Pt(10)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 4 — Technical architecture
    # ═════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4 · Technical Architecture Overview")

    arch = (
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│  DATA SOURCES          PROCESSING              OUTPUT            │\n"
        "│                                                                  │\n"
        "│  NIH Reporter API  →  5-Signal Scorer      →  Lead Radar Tab    │\n"
        "│  (90-day R01 window)   (AI · Data · Size ·     (ranked table +  │\n"
        "│                         Freshness · Type)        email + export) │\n"
        "│                                                                  │\n"
        "│  Grants.gov API    →  Sentence Embedding   →  Grant Advisor Tab │\n"
        "│  (open grants DB)      (all-MiniLM-L6-v2)      (top 5 matches   │\n"
        "│                              ↓                  + explanations)  │\n"
        "│                        Claude AI (Haiku)                         │\n"
        "│                        (LLM reranker +                           │\n"
        "│                         email inference)                         │\n"
        "└─────────────────────────────────────────────────────────────────┘"
    )

    arch_p = doc.add_paragraph()
    arch_p.paragraph_format.space_before = Pt(4)
    arch_p.paragraph_format.space_after  = Pt(6)
    arch_run = arch_p.add_run(arch)
    arch_run.font.name = "Courier New"
    arch_run.font.size = Pt(8)

    add_body(doc,
        "Stack: Python · Streamlit · Anthropic Claude · Sentence-Transformers · "
        "NIH Reporter API · Grants.gov · pandas · SMTP validation",
        space_after=4)

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
